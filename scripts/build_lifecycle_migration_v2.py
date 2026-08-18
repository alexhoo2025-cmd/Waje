#!/usr/bin/env python3
"""Build the Lifecycle Joint migration specification in HTML, Markdown, PNG and DOCX."""

from __future__ import annotations

import base64
import html
import json
import math
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SPEC_PATH = ROOT / "config/lifecycle_joint_report_spec.json"
KNOWLEDGE = ROOT / "knowledge/02-数据"
ASSET_DIR = KNOWLEDGE / "assets/lifecycle-joint-migration-v2"
MAIN_MD = KNOWLEDGE / "GM-Lifecycle-Pool-v2-Joint迁移至起源平台实施方案-2026-08-14.md"
FEISHU_MD = KNOWLEDGE / "GM-Lifecycle-Pool-v2-Joint迁移至起源平台实施方案-飞书存档版-2026-08-14.md"
HTML_PATH = KNOWLEDGE / "GM-Lifecycle-Pool-v2-Joint迁移至起源平台实施方案-2026-08-14.html"
DOCX_DIR = ROOT / "output/feishu"
DOCX_PATH = DOCX_DIR / "GM-Lifecycle-Pool-v2-Joint迁移至起源平台实施方案-飞书存档版-2026-08-14.docx"
TABLE_HELPER = Path(
    "/Users/robin/.codex/plugins/cache/openai-primary-runtime/documents/26.813.12317/skills/documents/scripts"
)
sys.path.insert(0, str(TABLE_HELPER))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


COLORS = {
    "navy": "14213D",
    "blue": "2667FF",
    "teal": "19A974",
    "amber": "F3B700",
    "red": "E45858",
    "ink": "17213A",
    "muted": "667085",
    "line": "D9E2F1",
    "paper": "FFFFFF",
    "canvas": "F4F7FB",
    "pale_blue": "EAF1FF",
    "pale_teal": "E8F7F1",
    "pale_amber": "FFF7DC",
}

FONT_MEDIUM = "/System/Library/Fonts/STHeiti Medium.ttc"
FONT_LIGHT = "/System/Library/Fonts/STHeiti Light.ttc"
DOCX_FONT = "Arial Unicode MS"


def font(size: int, medium: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_MEDIUM if medium else FONT_LIGHT, size=size)


def rgb(hex_value: str) -> tuple[int, int, int]:
    return tuple(int(hex_value[i : i + 2], 16) for i in (0, 2, 4))


def draw_text(draw: ImageDraw.ImageDraw, xy, value: str, size: int, color: str, medium=False, anchor=None):
    draw.text(xy, value, font=font(size, medium), fill=rgb(color), anchor=anchor)


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill), outline=rgb(outline) if outline else None, width=width)


def fit_text(value: str, max_chars: int) -> str:
    return value if len(value) <= max_chars else value[: max_chars - 1] + "…"


def mockup_base(report: dict) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    im = Image.new("RGB", (1600, 900), rgb(COLORS["canvas"]))
    draw = ImageDraw.Draw(im)
    draw.rectangle((0, 0, 1600, 78), fill=rgb(COLORS["navy"]))
    draw_text(draw, (54, 38), f"{report['no']:02d}  {report['name']}", 30, "FFFFFF", True, "lm")
    draw_text(draw, (1548, 29), "页面结构示意 · 数值仅作版式说明", 16, "B8C4DB", False, "ra")

    rounded(draw, (40, 98, 1560, 194), 18, "FFFFFF", COLORS["line"])
    filters = ["开始日期", "结束日期", "软件版本", "包体", "分包渠道"] + report.get("specific_filters", [])[:1]
    x = 64
    widths = [178, 178, 178, 178, 178, 178]
    for idx, label in enumerate(filters[:6]):
        w = widths[idx]
        draw_text(draw, (x, 119), label, 15, COLORS["muted"], True)
        rounded(draw, (x, 142, x + w, 178), 8, "F8FAFD", COLORS["line"])
        sample = "2026-07-27" if label == "开始日期" else "2026-08-10" if label == "结束日期" else "全部"
        draw_text(draw, (x + 12, 160), sample, 16, COLORS["ink"], False, "lm")
        x += w + 16
    rounded(draw, (1376, 132, 1534, 180), 12, COLORS["blue"])
    draw_text(draw, (1455, 156), "查询", 20, "FFFFFF", True, "mm")
    return im, draw


def draw_kpis(draw: ImageDraw.ImageDraw, report: dict):
    kpis = report["kpis"][:6]
    gap = 14
    total_w = 1520
    card_w = (total_w - gap * (len(kpis) - 1)) / len(kpis)
    for i, label in enumerate(kpis):
        x1 = 40 + i * (card_w + gap)
        x2 = x1 + card_w
        rounded(draw, (x1, 214, x2, 322), 16, "FFFFFF", COLORS["line"])
        accent = [COLORS["blue"], COLORS["teal"], COLORS["amber"], COLORS["blue"], COLORS["red"], COLORS["teal"]][i]
        draw.rectangle((x1, 214, x1 + 6, 322), fill=rgb(accent))
        draw_text(draw, (x1 + 20, 239), fit_text(label, 12), 15, COLORS["muted"], True)
        draw_text(draw, (x1 + 20, 280), sample_value(label), 25, COLORS["ink"], True)
        draw_text(draw, (x1 + 20, 305), "查询期汇总", 13, COLORS["muted"])


def draw_axes(draw, box, title):
    x1, y1, x2, y2 = box
    rounded(draw, box, 18, "FFFFFF", COLORS["line"])
    draw_text(draw, (x1 + 24, y1 + 26), title, 20, COLORS["ink"], True)
    plot = (x1 + 62, y1 + 74, x2 - 28, y2 - 48)
    px1, py1, px2, py2 = plot
    for i in range(5):
        y = py1 + i * (py2 - py1) / 4
        draw.line((px1, y, px2, y), fill=rgb("E9EEF6"), width=2)
    draw.line((px1, py1, px1, py2), fill=rgb("B6C1D4"), width=2)
    draw.line((px1, py2, px2, py2), fill=rgb("B6C1D4"), width=2)
    return plot


def draw_summary_chart(draw):
    plot = draw_axes(draw, (40, 344, 1078, 688), "查询期下注规模与RTP趋势")
    x1, y1, x2, y2 = plot
    heights = [0.55, 0.73, 0.61, 0.82, 0.67, 0.88, 0.76]
    line = [0.50, 0.57, 0.53, 0.62, 0.58, 0.66, 0.64]
    points = []
    step = (x2 - x1) / len(heights)
    for i, h in enumerate(heights):
        cx = x1 + step * (i + 0.5)
        bw = 24
        draw.rounded_rectangle((cx - bw - 3, y2 - h * 200, cx - 3, y2), radius=5, fill=rgb(COLORS["blue"]))
        draw.rounded_rectangle((cx + 3, y2 - h * 218, cx + bw + 3, y2), radius=5, fill=rgb("83A8FF"))
        py = y2 - line[i] * 230
        points.append((cx, py))
        draw_text(draw, (cx, y2 + 22), f"D{i+1}", 13, COLORS["muted"], anchor="mm")
    draw.line(points, fill=rgb(COLORS["teal"]), width=5, joint="curve")
    for p in points:
        draw.ellipse((p[0]-6, p[1]-6, p[0]+6, p[1]+6), fill=rgb(COLORS["teal"]), outline=rgb("FFFFFF"), width=2)


def draw_detail_chart(draw):
    rounded(draw, (40, 344, 1078, 688), 18, "FFFFFF", COLORS["line"])
    draw_text(draw, (64, 370), "生命周期 × 游戏：下注贡献与RTP偏差矩阵", 20, COLORS["ink"], True)
    games = ["Whot", "Fish", "Tada", "Soccer", "Roulette", "Slots"]
    lifes = ["L1", "L2", "L3", "L4"]
    values = [[0.2,0.5,0.8,0.35,0.62,0.45],[0.4,0.7,0.3,0.58,0.85,0.52],[0.68,0.45,0.74,0.26,0.55,0.9],[0.33,0.64,0.48,0.8,0.42,0.6]]
    sx, sy, cw, ch = 152, 432, 126, 48
    for j, g in enumerate(games):
        draw_text(draw, (sx+j*cw+cw/2, 414), g, 14, COLORS["muted"], True, "mm")
    for i, life in enumerate(lifes):
        draw_text(draw, (112, sy+i*(ch+10)+ch/2), life, 15, COLORS["ink"], True, "mm")
        for j, v in enumerate(values[i]):
            tone = (int(238-80*v), int(246-25*v), int(255-8*v))
            draw.rounded_rectangle((sx+j*cw, sy+i*(ch+10), sx+j*cw+112, sy+i*(ch+10)+ch), radius=8, fill=tone)
            draw_text(draw, (sx+j*cw+56, sy+i*(ch+10)+ch/2), f"{(v-.5)*8:+.1f}pp", 14, COLORS["ink"], True, "mm")


def draw_game_chart(draw):
    plot = draw_axes(draw, (40, 344, 1078, 688), "游戏下注贡献与实际 / 预期RTP")
    x1, y1, x2, y2 = plot
    games = ["Whot", "Fish", "Tada", "Soccer", "Roulette"]
    shares = [0.92,0.76,0.62,0.47,0.31]
    for i, (g, s) in enumerate(zip(games, shares)):
        y = y1 + 28 + i*48
        draw_text(draw, (x1-12, y+10), g, 14, COLORS["muted"], True, "ra")
        draw.rounded_rectangle((x1, y, x1+(x2-x1)*s*0.78, y+20), radius=8, fill=rgb(COLORS["blue"]))
        expected_x = x1+(x2-x1)*(0.55+0.06*i)
        actual_x = expected_x + [34,-28,17,-42,25][i]
        draw.line((expected_x, y-4, expected_x, y+25), fill=rgb(COLORS["amber"]), width=4)
        draw.ellipse((actual_x-6,y+4,actual_x+6,y+16), fill=rgb(COLORS["teal"]))


def draw_return_chart(draw):
    plot = draw_axes(draw, (40, 344, 1078, 688), "生命周期基础 / 完全下注与RTP")
    x1, y1, x2, y2 = plot
    base = [0.34,0.53,0.82,0.66]
    entire = [0.38,0.58,0.88,0.73]
    line = [0.57,0.66,0.61,0.74]
    step=(x2-x1)/4
    points=[]
    for i in range(4):
        cx=x1+step*(i+.5)
        draw.rounded_rectangle((cx-34,y2-base[i]*230,cx-5,y2),radius=5,fill=rgb(COLORS["blue"]))
        draw.rounded_rectangle((cx+5,y2-entire[i]*230,cx+34,y2),radius=5,fill=rgb("83A8FF"))
        points.append((cx,y2-line[i]*230))
        draw_text(draw,(cx,y2+22),f"生命周期{i+1}",13,COLORS["muted"],True,"mm")
    draw.line(points,fill=rgb(COLORS["teal"]),width=5)
    for p in points: draw.ellipse((p[0]-6,p[1]-6,p[0]+6,p[1]+6),fill=rgb(COLORS["teal"]))


def draw_payment_chart(draw):
    plot = draw_axes(draw, (40, 344, 1078, 688), "生命周期充值 / 营收 / TX 与TC比")
    x1,y1,x2,y2=plot
    recharge=[.42,.63,.86,.71]; revenue=[.28,.49,.72,.58]; tx=[.18,.32,.57,.44]; tc=[.48,.57,.65,.60]
    step=(x2-x1)/4; points=[]
    for i in range(4):
        cx=x1+step*(i+.5)
        for off,val,col in [(-38,recharge[i],COLORS["blue"]),(-10,revenue[i],COLORS["teal"]),(18,tx[i],COLORS["amber"])]:
            draw.rounded_rectangle((cx+off,y2-val*230,cx+off+24,y2),radius=4,fill=rgb(col))
        points.append((cx,y2-tc[i]*230))
        draw_text(draw,(cx,y2+22),f"生命周期{i+1}",13,COLORS["muted"],True,"mm")
    draw.line(points,fill=rgb(COLORS["red"]),width=5)
    for p in points: draw.ellipse((p[0]-6,p[1]-6,p[0]+6,p[1]+6),fill=rgb(COLORS["red"]))


def draw_side_panel(draw, report):
    rounded(draw, (1098, 344, 1560, 688), 18, "FFFFFF", COLORS["line"])
    draw_text(draw, (1124, 374), "查询期汇总规则", 20, COLORS["ink"], True)
    items = [
        ("分组", report["group_by"]),
        ("金额", "查询期直接求和"),
        ("比例", "累计分子/分母重算"),
        ("人数", "按原用户键查询期去重"),
        ("空值", "零分母显示 N/A"),
    ]
    y=418
    for label,value in items:
        rounded(draw,(1122,y,1184,y+28),8,COLORS["pale_blue"])
        draw_text(draw,(1153,y+14),label,13,COLORS["blue"],True,"mm")
        lines=[]
        text=value
        while text:
            lines.append(text[:18]); text=text[18:]
        for line_idx,line in enumerate(lines[:2]):
            draw_text(draw,(1198,y+4+line_idx*20),line,14,COLORS["ink"])
        y+=55


def draw_table_preview(draw, report):
    rounded(draw,(40,712,1560,858),18,"FFFFFF",COLORS["line"])
    draw_text(draw,(64,740),"结果表（查询期汇总）",18,COLORS["ink"],True)
    fields=[f["label"] for f in report["fields"] if f["key"]!="stat_date"][:7]
    x=64; y=770; available=1468; w=available/max(len(fields),1)
    for i,label in enumerate(fields):
        x1=x+i*w
        draw.rectangle((x1,y,x1+w,y+34),fill=rgb(COLORS["pale_blue"]),outline=rgb(COLORS["line"]))
        draw_text(draw,(x1+w/2,y+17),fit_text(label,9),13,COLORS["ink"],True,"mm")
        draw.rectangle((x1,y+34,x1+w,y+74),fill=rgb("FFFFFF"),outline=rgb(COLORS["line"]))
        sample=sample_value(label)
        draw_text(draw,(x1+w/2,y+54),sample,13,COLORS["muted"],False,"mm")


def sample_value(label: str) -> str:
    if "最大游戏" in label:
        return "Whot +2.4pp"
    if label == "游戏数":
        return "28"
    if "生命周期" in label:
        return "1"
    if label in {"游戏", "游戏类型"}:
        return "Whot"
    if "人数" in label or "用户数" in label:
        return "53,800"
    if "次数" in label:
        return "8,063"
    if "回报比" in label or "TC比" in label or "流充比" in label or "占比" in label or "系数" in label:
        return "96.47%"
    if "人均" in label:
        return "1,044.32"
    if "下注额" in label:
        return "1.53B"
    if "盈利" in label or "金额" in label or "营收" in label or "TX" in label:
        return "56.12M"
    return "—"


def build_mockups(spec: dict) -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    outputs={}
    chart_funcs={1:draw_summary_chart,2:draw_detail_chart,3:draw_game_chart,4:draw_return_chart,5:draw_payment_chart}
    for report in spec["reports"]:
        im,draw=mockup_base(report)
        draw_kpis(draw,report)
        chart_funcs[report["no"]](draw)
        draw_side_panel(draw,report)
        draw_table_preview(draw,report)
        path=ASSET_DIR/f"{report['no']:02d}-{report['id']}.png"
        im.save(path,optimize=True)
        outputs[report["id"]]=path
    return outputs


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    lines=["| " + " | ".join(headers) + " |", "|" + "|".join(["---"]*len(headers)) + "|"]
    for row in rows:
        lines.append("| " + " | ".join(str(v).replace("|","｜") for v in row) + " |")
    return "\n".join(lines)


def report_filters(spec, report):
    return [f["label"] for f in spec["global_filters"]] + report.get("specific_filters",[])


def build_markdown(spec: dict, assets: dict[str,Path], feishu=False) -> str:
    lines=[]
    if not feishu:
        lines += ["---",f"title: {spec['title']}",f"date: {spec['date']}","type: data-platform-migration-spec",f"status: {spec['status']}","---",""]
    lines += [f"# {spec['title']}（{spec['version']}）","",f"> {spec['decision']}",""]
    lines += ["## 一页结论","",md_table(["项目","确定方案"],[
        ["报表数量","5张独立报表"],
        ["默认统计","查询期汇总；可切换按日明细"],
        ["必备筛选","日期范围、软件版本、包体、分包渠道、归因媒体、归因渠道"],
        ["口径","沿用GM原指标算法；比例与人均从累计分子/分母重算"],
        ["权限","起源只读；GM修改能力不迁移"],
        ["触发","仅按需手动触发；收到日期范围与更新要求后由用户点击查询，不创建定时任务"],
    ]),""]
    lines += ["## 五张报表","",md_table(["序号","报表名称","查询期分组","核心用途"],[[r["no"],r["name"],r["group_by"],r["purpose"]] for r in spec["reports"]]),""]
    lines += ["## 共用筛选字段","",md_table(["筛选字段","控件","默认值","说明"],[[f["label"],f["control"],f["default"],"必填" if f["required"] else "可选"] for f in spec["global_filters"]]),""]
    lines += ["## 查询期汇总原则","",md_table(["类型","规则"],[[r["label"],r["rule"]] for r in spec["aggregation_rules"]]),""]
    for report in spec["reports"]:
        lines += [f"## 报表 {report['no']}｜{report['name']}","",report["purpose"],"",f"**筛选字段：** {'、'.join(report_filters(spec,report))}","",f"**汇总粒度：** {report['group_by']}","",f"**图表：** {report['chart']}",""]
        rel=assets[report["id"]].relative_to(KNOWLEDGE)
        lines += [f"![{report['name']}页面示意]({rel.as_posix()})","","> 图为页面结构示意，数字不作为业务结论。",""]
        rows=[]
        rules={r["type"]:r["label"] for r in spec["aggregation_rules"]}
        for f in report["fields"]:
            rows.append([f["label"],f["key"],rules[f["aggregation"]],f["note"]])
        lines += [md_table(["显示字段","数据键","汇总类型","查询期算法"],rows),""]
    lines += ["## 实施规则","", *[f"- {x}" for x in spec["implementation_rules"]],"", "## 验收清单","", *[f"- {x}" for x in spec["acceptance"]],""]
    lines += ["## 现场依据","",f"- [旧GM Joint报表]({spec['source_url']})",f"- [起源分析平台]({spec['origin_url']})",f"- [最终在线结果文档]({spec['lark_url']})",""]
    if feishu:
        lines += ["> 存档说明：本版本按飞书连续阅读习惯缩短段落、拆分宽表，并保留5张页面示意图。"]
    return "\n".join(lines)


def image_data_url(path: Path) -> str:
    return "data:image/png;base64,"+base64.b64encode(path.read_bytes()).decode("ascii")


def build_html(spec: dict, assets: dict[str,Path]) -> str:
    report_cards="".join(f"<a href='#r{r['no']}'><b>{r['no']:02d}</b><span>{html.escape(r['name'])}</span><small>{html.escape(r['group_by'])}</small></a>" for r in spec["reports"])
    global_filters="".join(f"<span>{html.escape(f['label'])}</span>" for f in spec["global_filters"])
    sections=[]
    rules={r["type"]:r["label"] for r in spec["aggregation_rules"]}
    for r in spec["reports"]:
        filters="".join(f"<span>{html.escape(x)}</span>" for x in report_filters(spec,r))
        rows="".join(f"<tr><td>{html.escape(f['label'])}</td><td><code>{html.escape(f['key'])}</code></td><td>{html.escape(rules[f['aggregation']])}</td><td>{html.escape(f['note'])}</td></tr>" for f in r["fields"])
        kpis="".join(f"<span>{html.escape(x)}</span>" for x in r["kpis"])
        sections.append(f"""
        <section class="report" id="r{r['no']}">
          <div class="section-head"><div class="num">{r['no']:02d}</div><div><h2>{html.escape(r['name'])}</h2><p>{html.escape(r['purpose'])}</p></div></div>
          <div class="two"><div><h3>筛选字段</h3><div class="chips">{filters}</div></div><div><h3>查询期汇总</h3><p>{html.escape(r['group_by'])}</p></div></div>
          <div class="kpi-line"><b>顶部指标</b>{kpis}</div>
          <figure><img src="{image_data_url(assets[r['id']])}" alt="{html.escape(r['name'])}页面结构示意"><figcaption>{html.escape(r['chart'])}；示意数值不作为业务结论。</figcaption></figure>
          <h3>显示字段与区间算法</h3>
          <div class="table-wrap"><table><thead><tr><th>显示字段</th><th>数据键</th><th>汇总类型</th><th>查询期算法</th></tr></thead><tbody>{rows}</tbody></table></div>
        </section>""")
    agg="".join(f"<div><b>{html.escape(r['label'])}</b><p>{html.escape(r['rule'])}</p></div>" for r in spec["aggregation_rules"])
    impl="".join(f"<li>{html.escape(x)}</li>" for x in spec["implementation_rules"])
    acceptance="".join(f"<li>{html.escape(x)}</li>" for x in spec["acceptance"])
    return f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{html.escape(spec['title'])}（{spec['version']}）</title>
<style>
:root{{--navy:#14213d;--blue:#2667ff;--teal:#19a974;--amber:#f3b700;--ink:#17213a;--muted:#667085;--line:#d9e2f1;--canvas:#f4f7fb;--paper:#fff}}
*{{box-sizing:border-box}}html{{scroll-behavior:smooth}}body{{margin:0;background:var(--canvas);color:var(--ink);font:15px/1.65 -apple-system,BlinkMacSystemFont,"PingFang SC","Microsoft YaHei",Arial,sans-serif}}
.hero{{background:linear-gradient(135deg,#14213d,#233b70);color:#fff;padding:70px 24px 58px}}.hero>div,.main{{max-width:1180px;margin:auto}}.eyebrow{{letter-spacing:.16em;text-transform:uppercase;color:#a9c0ff;font-size:12px;font-weight:800}}h1{{font-size:clamp(34px,5vw,58px);line-height:1.08;letter-spacing:-.04em;margin:14px 0 18px;max-width:950px}}.lead{{max-width:900px;font-size:19px;color:#dbe4f6}}.meta{{display:flex;gap:10px;flex-wrap:wrap;margin-top:24px}}.meta span{{border:1px solid #52648a;border-radius:999px;padding:7px 12px;color:#e8edfa}}
.main{{padding:34px 20px 80px}}.notice{{background:#fff7dc;border:1px solid #f2d77c;border-radius:18px;padding:18px 22px;margin-bottom:24px}}.notice b{{color:#825e00}}h2{{font-size:30px;letter-spacing:-.025em;margin:0}}h3{{font-size:18px;margin:0 0 10px}}p{{margin:5px 0 12px;color:#465169}}
.map{{display:grid;grid-template-columns:repeat(5,1fr);gap:12px;margin:22px 0 34px}}.map a{{display:flex;min-height:132px;flex-direction:column;padding:18px;border:1px solid var(--line);border-radius:16px;background:#fff;color:var(--ink);text-decoration:none;box-shadow:0 8px 26px rgba(18,33,61,.05)}}.map b{{color:var(--blue);font-size:20px}}.map span{{font-weight:800;margin:8px 0}}.map small{{color:var(--muted)}}
.common,.report{{background:#fff;border:1px solid var(--line);border-radius:22px;padding:28px;margin:24px 0;box-shadow:0 12px 36px rgba(18,33,61,.06)}}.filters,.chips{{display:flex;gap:8px;flex-wrap:wrap}}.filters span,.chips span,.kpi-line span{{padding:6px 10px;background:#eaf1ff;color:#174cb4;border-radius:999px;font-weight:700;font-size:13px}}.agg-grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:12px;margin-top:18px}}.agg-grid div{{background:#f7f9fc;border:1px solid var(--line);border-radius:14px;padding:15px}}.agg-grid p{{font-size:13px;margin:4px 0}}
.section-head{{display:flex;gap:16px;align-items:center}}.num{{width:58px;height:58px;border-radius:16px;background:var(--navy);color:#fff;display:grid;place-items:center;font-size:22px;font-weight:900}}.two{{display:grid;grid-template-columns:1.35fr 1fr;gap:16px;margin:20px 0}}.two>div{{border:1px solid var(--line);border-radius:14px;padding:16px}}.kpi-line{{display:flex;align-items:center;gap:8px;flex-wrap:wrap;margin:16px 0}}.kpi-line b{{margin-right:4px}}figure{{margin:18px 0}}figure img{{width:100%;display:block;border:1px solid var(--line);border-radius:16px}}figcaption{{color:var(--muted);font-size:13px;margin-top:7px}}
.table-wrap{{overflow:auto;border:1px solid var(--line);border-radius:14px}}table{{width:100%;border-collapse:collapse;min-width:760px}}th,td{{padding:11px 12px;border-bottom:1px solid var(--line);text-align:left;vertical-align:top}}th{{background:#eaf1ff;color:#173d88;white-space:nowrap}}tr:last-child td{{border-bottom:0}}code{{font-size:12px;background:#eef2f7;padding:2px 5px;border-radius:5px}}.check-grid{{display:grid;grid-template-columns:1fr 1fr;gap:18px}}ol,ul{{padding-left:22px}}li{{margin:7px 0}}
.footer{{max-width:1180px;margin:auto;padding:24px 20px 42px;color:var(--muted);border-top:1px solid var(--line)}}
@media(max-width:900px){{.map{{grid-template-columns:1fr 1fr}}.agg-grid,.two,.check-grid{{grid-template-columns:1fr}}.common,.report{{padding:20px}}}}@media(max-width:560px){{.map{{grid-template-columns:1fr}}h1{{font-size:36px}}}}
@media print{{body{{background:#fff}}.hero{{padding:34px 0}}.main{{padding:20px 0}}.common,.report{{box-shadow:none;break-inside:avoid}}}}
</style></head><body>
<header class="hero"><div><div class="eyebrow">WAJE DATA PLATFORM · IMPLEMENTATION SPEC</div><h1>{html.escape(spec['title'])}</h1><p class="lead">5张独立报表、查询期汇总、版本/包体/渠道筛选，以及与最终在线结果文档一致的字段合同。</p><div class="meta"><span>{spec['version']}</span><span>{spec['date']}</span><span>{html.escape(spec['status'])}</span></div></div></header>
<main class="main"><div class="notice"><b>确定方案：</b>{html.escape(spec['decision'])}</div><div class="map">{report_cards}</div>
<section class="common"><h2>共用筛选</h2><p>五张报表复用同一套筛选字典。仅在用户给出日期范围与更新要求后手动触发，并由用户点击“查询”；不配置定时任务或后台自动调度。</p><div class="filters">{global_filters}</div><div class="agg-grid">{agg}</div></section>
{''.join(sections)}
<section class="common"><div class="check-grid"><div><h2>实施规则</h2><ul>{impl}</ul></div><div><h2>验收清单</h2><ol>{acceptance}</ol></div></div></section>
</main><footer class="footer">现场依据：旧GM Joint报表、起源分析平台、最终在线结果文档。所有页面图示仅表达版式，不代表真实业务数值。</footer></body></html>"""


def set_repeat_table_header(row):
    tr_pr=row._tr.get_or_add_trPr(); el=OxmlElement("w:tblHeader"); el.set(qn("w:val"),"true"); tr_pr.append(el)


def set_cell_shading(cell, fill):
    tc_pr=cell._tc.get_or_add_tcPr(); shd=tc_pr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"),fill)


def set_cell_text(cell, text, *, bold=False, color="17213A", size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text=""; p=cell.paragraphs[0]; p.alignment=align; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
    run=p.add_run(str(text)); run.bold=bold; run.font.name=DOCX_FONT; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),DOCX_FONT); run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, weights, font_size=8.5):
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"
    for i,h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i],h,bold=True,color="173D88",size=9,align=WD_ALIGN_PARAGRAPH.CENTER); set_cell_shading(table.rows[0].cells[i],COLORS["pale_blue"])
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells=table.add_row().cells
        for i,value in enumerate(row):
            set_cell_text(cells[i],value,size=font_size,align=WD_ALIGN_PARAGRAPH.CENTER if i==0 else WD_ALIGN_PARAGRAPH.LEFT)
    widths=column_widths_from_weights(weights,9360); apply_table_geometry(table,widths,table_width_dxa=9360,indent_dxa=120,cell_margins_dxa={"top":90,"bottom":90,"start":120,"end":120})
    for row in table.rows:
        row.height_rule=WD_ROW_HEIGHT_RULE.AT_LEAST
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return table


def set_run(run, size=11, color="17213A", bold=False, italic=False):
    run.font.name=DOCX_FONT; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),DOCX_FONT); run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold; run.italic=italic


def add_para(doc,text,size=11,color="17213A",bold=False,italic=False,after=6,align=WD_ALIGN_PARAGRAPH.LEFT,style=None):
    p=doc.add_paragraph(style=style); p.alignment=align; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.25
    set_run(p.add_run(text),size,color,bold,italic); return p


def add_callout(doc,text,fill="FFF7DC",color="6F5200"):
    table=doc.add_table(rows=1,cols=1); set_cell_text(table.cell(0,0),text,bold=True,color=color,size=10.5); set_cell_shading(table.cell(0,0),fill); set_repeat_table_header(table.rows[0]); apply_table_geometry(table,[9360],table_width_dxa=9360,indent_dxa=190,cell_margins_dxa={"top":160,"bottom":160,"start":190,"end":190}); doc.add_paragraph().paragraph_format.space_after=Pt(2)


def add_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run=paragraph.add_run(); fld_char1=OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"),"begin"); instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=" PAGE "; fld_char2=OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"),"end"); run._r.extend([fld_char1,instr,fld_char2]); set_run(run,9,"667085")


def style_doc(doc):
    section=doc.sections[0]; section.page_width=Inches(8.5); section.page_height=Inches(11); section.top_margin=Inches(1); section.bottom_margin=Inches(1); section.left_margin=Inches(1); section.right_margin=Inches(1); section.header_distance=Inches(.492); section.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name=DOCX_FONT; normal._element.rPr.rFonts.set(qn("w:eastAsia"),DOCX_FONT); normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(COLORS["ink"]); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,color,before,after in [("Heading 1",16,"2E74B5",18,10),("Heading 2",13,"2E74B5",14,7),("Heading 3",12,"1F4D78",10,5)]:
        st=styles[name]; st.font.name=DOCX_FONT; st._element.rPr.rFonts.set(qn("w:eastAsia"),DOCX_FONT); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    header=section.header; hp=header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT; hp.paragraph_format.space_after=Pt(0); set_run(hp.add_run("WAJE 数据平台迁移方案  |  Lifecycle Pool v2 (Joint)"),8.5,"667085",True)
    footer=section.footer; fp=footer.paragraphs[0]; add_page_number(fp)


def build_docx(spec,assets):
    DOCX_DIR.mkdir(parents=True,exist_ok=True)
    doc=Document(); style_doc(doc)
    add_para(doc,"WAJE DATA PLATFORM",10,COLORS["blue"],True,after=8)
    add_para(doc,spec["title"],24,COLORS["navy"],True,after=8)
    add_para(doc,"5张独立报表 · 查询期汇总 · 版本 / 包体 / 渠道筛选",13,COLORS["muted"],after=18)
    add_table(doc,["版本","日期","状态"],[[spec["version"],spec["date"],spec["status"]]],[1,1.2,2.6],9)
    add_callout(doc,"确定方案："+spec["decision"])
    doc.add_heading("一页结论",level=1)
    add_table(doc,["报表","查询期分组","核心用途"],[[r["name"],r["group_by"],r["purpose"]] for r in spec["reports"]],[1.8,2.35,2.35],8.5)
    doc.add_heading("共用筛选",level=1)
    add_table(doc,["筛选字段","控件","默认值"],[[f["label"],f["control"],f["default"]] for f in spec["global_filters"]],[1.5,1.25,3.75],9)
    doc.add_heading("查询期汇总原则",level=1)
    add_table(doc,["类型","规则"],[[r["label"],r["rule"]] for r in spec["aggregation_rules"]],[1.35,5.15],9)
    rules={r["type"]:r["label"] for r in spec["aggregation_rules"]}
    for report in spec["reports"]:
        doc.add_page_break(); doc.add_heading(f"报表 {report['no']}｜{report['name']}",level=1)
        add_para(doc,report["purpose"],11.5,COLORS["ink"],True,after=8)
        add_para(doc,"筛选字段："+"、".join(report_filters(spec,report)),10,COLORS["muted"],after=4)
        add_para(doc,"汇总粒度："+report["group_by"],10,COLORS["muted"],after=8)
        picture=doc.add_picture(str(assets[report["id"]]),width=Inches(6.25))
        picture._inline.docPr.set("title",report["name"]+"页面结构示意")
        picture._inline.docPr.set("descr",report["chart"]+"；页面结构示意，图中数值非真实业务数据。")
        cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after=Pt(10); set_run(cap.add_run(report["chart"]+"（页面结构示意，数值非真实数据）"),8.5,COLORS["muted"],False,True)
        doc.add_heading("显示字段与查询期算法",level=2)
        rows=[[f["label"],f["key"],rules[f["aggregation"]]+"；"+f["note"]] for f in report["fields"]]
        add_table(doc,["显示字段","数据键","查询期算法"],rows,[1.65,2.3,2.55],8)
    doc.add_page_break(); doc.add_heading("实施规则",level=1)
    for idx,item in enumerate(spec["implementation_rules"],1): add_para(doc,f"{idx}. {item}",9.5,after=2)
    doc.add_heading("验收清单",level=1)
    for idx,item in enumerate(spec["acceptance"],1): add_para(doc,f"{idx}. {item}",9.5,after=2)
    doc.add_heading("现场依据",level=1)
    for label,url in [("旧GM Joint报表",spec["source_url"]),("起源分析平台",spec["origin_url"]),("最终在线结果文档",spec["lark_url"])]: add_para(doc,f"{label}：{url}",8.5,COLORS["muted"],after=2)
    doc.core_properties.title=spec["title"]+"（飞书存档版）"; doc.core_properties.subject="5张报表迁移、筛选与汇总算法"; doc.core_properties.author="Waje Data Product"
    doc.save(DOCX_PATH)


def main():
    spec=json.loads(SPEC_PATH.read_text(encoding="utf-8")); assets=build_mockups(spec)
    MAIN_MD.write_text(build_markdown(spec,assets,False),encoding="utf-8")
    FEISHU_MD.write_text(build_markdown(spec,assets,True),encoding="utf-8")
    HTML_PATH.write_text(build_html(spec,assets),encoding="utf-8")
    build_docx(spec,assets)
    print(json.dumps({"html":str(HTML_PATH),"markdown":str(MAIN_MD),"feishu_markdown":str(FEISHU_MD),"docx":str(DOCX_PATH),"assets":[str(p) for p in assets.values()]},ensure_ascii=False,indent=2))


if __name__=="__main__": main()
